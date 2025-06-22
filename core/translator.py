# core/translator.py

import os
import time
import re
import shutil
import google.generativeai as genai
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from docx import Document
from google.api_core.exceptions import ResourceExhausted

from .project_manager import PROJECTS_DIR, ProjectManager


def translation_process(project_data, progress_queue, stop_event):
    pm = ProjectManager()
    try:
        project_name = project_data["project_name"]
        completed_chapters_list = project_data["completed_chapters_list"]

        user_prompt = project_data["prompt"]
        if "{text_to_translate}" not in user_prompt:
            progress_queue.put(("log", "⚠️ Обнаружен старый формат промпта. Автоматически модернизируем его."))
            modern_prompt_template = (
                "{user_prompt_text}\n\n"
                "You are a professional literary translator. Translate the following text from English into Russian.\n"
                "Preserve the original style, tone, and formatting (paragraphs, line breaks). "
                "Translate the meaning accurately, not just word for word.\n"
                "{glossary}\n"
                "Text to translate:\n"
                "---\n"
                "{text_to_translate}"
            )
            final_prompt_template = modern_prompt_template.format(
                user_prompt_text=user_prompt,
                glossary="{glossary}",
                text_to_translate="{text_to_translate}"
            )
        else:
            final_prompt_template = user_prompt

        temp_dir = os.path.join(PROJECTS_DIR, project_name, "temp")
        if not project_data["resume"]:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            completed_chapters_list = []
        os.makedirs(temp_dir, exist_ok=True)

        genai.configure(api_key=project_data["api_key"])
        model = genai.GenerativeModel(project_data["model"])

        # 1. Парсим глоссарий из текстового поля
        glossary = {}
        for line in project_data["glossary"].split('\n'):
            if '->' in line and not line.strip().startswith('#'):
                parts = line.split('->', 1)
                original, translation = parts[0].strip(), parts[1].strip()
                if original and translation:
                    glossary[original] = translation

        # 2. Формируем инструкции для AI на основе глоссария
        glossary_instructions = ""
        if glossary:
            instructions_list = ["\nStrictly follow these translation rules:"]
            for original, translation in glossary.items():
                original_clean = original.strip("'\"")
                translation_clean = translation.strip("'\"")
                instructions_list.append(f'- Translate "{original_clean}" as "{translation_clean}".')
            glossary_instructions = "\n".join(instructions_list) + "\n"

        progress_queue.put(("log", f"Используется модель: {project_data['model']}"))
        book = epub.read_epub(project_data["epub_path"])
        items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT))
        total_items = len(items)

        for i, item in enumerate(items):
            if stop_event.is_set():
                break

            if i in completed_chapters_list:
                progress_queue.put(("log", f"Глава {i + 1} уже переведена. Пропускаем."))
                progress_queue.put(("progress", (i + 1, total_items)))
                continue

            progress_queue.put(("progress", (i, total_items)))
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            original_text = soup.get_text(separator='\n', strip=True)
            if not original_text.strip():
                progress_queue.put(("log", f"Глава {i + 1} пустая, пропускаем."))
                completed_chapters_list.append(i)
                pm.update_completed_chapters(project_name, completed_chapters_list)
                progress_queue.put(("progress", (i + 1, total_items)))
                continue

            # 3. Собираем финальный промпт, вставляя инструкции и текст для перевода
            # Используем `final_prompt_template`, который был подготовлен в начале функции
            prompt = final_prompt_template.format(
                glossary=glossary_instructions,
                text_to_translate=original_text
            )

            translated_text = ""
            max_retries = 5
            retry_delay = 10

            safety_settings = {
                "HARM_CATEGORY_HARASSMENT": "BLOCK_NONE",
                "HARM_CATEGORY_HATE_SPEECH": "BLOCK_NONE",
                "HARM_CATEGORY_SEXUALLY_EXPLICIT": "BLOCK_NONE",
                "HARM_CATEGORY_DANGEROUS_CONTENT": "BLOCK_NONE",
            }

            for attempt in range(max_retries):
                if stop_event.is_set():
                    break
                try:
                    progress_queue.put(
                        ("log", f"Глава {i + 1}: Отправка запроса в API (попытка {attempt + 1}/{max_retries})..."))

                    response = model.generate_content(prompt, safety_settings=safety_settings)

                    try:
                        translated_text = response.text
                    except ValueError:
                        finish_reason = "Неизвестно"
                        if response.prompt_feedback and response.prompt_feedback.block_reason:
                            finish_reason = f"Заблокировано по причине: {response.prompt_feedback.block_reason.name}"
                        elif response.candidates and response.candidates[0].finish_reason:
                            finish_reason = f"Причина завершения: {response.candidates[0].finish_reason.name} ({response.candidates[0].finish_reason.value})"

                        progress_queue.put(("log", f"⚠️ Глава {i + 1}: Ответ от API пустой. {finish_reason}"))
                        translated_text = ""

                    progress_queue.put(("log", f"Глава {i + 1}: Ответ от API получен."))
                    break

                except ResourceExhausted as e:
                    progress_queue.put((
                        "log",
                        f"⚠️ Превышен лимит API для главы {i + 1}. Попытка {attempt + 1}/{max_retries}. "
                        f"Ждем {retry_delay} секунд..."
                    ))
                    for _ in range(retry_delay):
                        if stop_event.is_set(): break
                        time.sleep(1)
                    if stop_event.is_set(): break
                    retry_delay *= 2

                except Exception as e:
                    progress_queue.put(("log", f"Критическая ошибка API: {e}"))
                    raise e

            if stop_event.is_set():
                break

            if not translated_text:
                progress_queue.put(("log",
                                    f"❌ Не удалось получить перевод для главы {i + 1} после {max_retries} попыток. Пропускаем."))
                continue

            temp_file_path = os.path.join(temp_dir, f"chapter_{i:04d}.txt")
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                chapter_title_tag = soup.find(['h1', 'h2', 'h3'])
                chapter_title = chapter_title_tag.get_text(strip=True) if chapter_title_tag else f"Глава {i + 1}"
                f.write(f"<h1>{chapter_title}</h1>\n{translated_text}")

            completed_chapters_list.append(i)
            pm.update_completed_chapters(project_name, completed_chapters_list)

            progress_queue.put(("progress", (i + 1, total_items)))
            if not stop_event.is_set() and project_data["delay"] > 0:
                progress_queue.put(("log", f"Задержка на {project_data['delay']} сек..."))
                time.sleep(project_data["delay"])

        if not stop_event.is_set():
            progress_queue.put(("log", "Все главы переведены. Собираем DOCX..."))
            doc = Document()
            metadata_title = book.get_metadata('DC', 'title')
            doc.add_heading(metadata_title[0][0] if metadata_title else "Переведенная книга", 0)

            if os.path.exists(temp_dir) and os.listdir(temp_dir):
                for filename in sorted(os.listdir(temp_dir)):
                    if filename.endswith(".txt"):
                        with open(os.path.join(temp_dir, filename), 'r', encoding='utf-8') as f:
                            content = f.read().split('\n', 1)
                            title = content[0].replace("<h1>", "").replace("</h1>", "")
                            text = content[1] if len(content) > 1 else ""
                            doc.add_heading(title, level=1)
                            doc.add_paragraph(text)
                            doc.add_page_break()

            doc.save(project_data["output_path"])
            pm.cleanup_project(project_name)
            progress_queue.put(("done", None))
        else:
            progress_queue.put(("log", "Перевод отменен. Прогресс сохранен."))

    except Exception as e:
        import traceback
        progress_queue.put(("error", traceback.format_exc()))
    finally:
        progress_queue.put(("finish_signal", None))